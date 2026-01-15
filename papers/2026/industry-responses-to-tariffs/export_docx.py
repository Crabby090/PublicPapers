"""
Export paper.md to formatted DOCX for journal submission.
Times New Roman, double spacing, proper heading styles.
"""

import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path

def set_style_font(style, font_name='Times New Roman', font_size=12):
    """Set font for a style."""
    style.font.name = font_name
    style.font.size = Pt(font_size)

def create_formatted_docx(md_path, output_path):
    """Convert markdown to formatted DOCX."""

    # Read markdown
    with open(md_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Create document
    doc = Document()

    # Set up styles
    styles = doc.styles

    # Normal style - Times New Roman 12pt, double spaced
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)

    # Heading 1 - Title
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h1_style.paragraph_format.space_before = Pt(0)
    h1_style.paragraph_format.space_after = Pt(0)
    h1_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 2 - Section headers
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(0)
    h2_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 3 - Subsection headers
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(0)
    h3_style.paragraph_format.first_line_indent = Inches(0)

    # Quote style
    try:
        quote_style = styles.add_style('Block Quote', WD_STYLE_TYPE.PARAGRAPH)
    except:
        quote_style = styles['Block Quote']
    quote_style.font.name = 'Times New Roman'
    quote_style.font.size = Pt(11)
    quote_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.right_indent = Inches(0.5)
    quote_style.paragraph_format.first_line_indent = Inches(0)

    # Reference style (hanging indent)
    try:
        ref_style = styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
    except:
        ref_style = styles['Reference']
    ref_style.font.name = 'Times New Roman'
    ref_style.font.size = Pt(12)
    ref_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    ref_style.paragraph_format.left_indent = Inches(0.5)
    ref_style.paragraph_format.first_line_indent = Inches(-0.5)
    ref_style.paragraph_format.space_after = Pt(0)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Parse and convert markdown
    lines = content.split('\n')
    i = 0
    in_blockquote = False
    blockquote_text = []

    while i < len(lines):
        line = lines[i]

        # Skip horizontal rules
        if line.strip() == '---':
            i += 1
            continue

        # Handle blockquotes
        if line.startswith('>'):
            quote_line = line[1:].strip()
            blockquote_text.append(quote_line)
            i += 1
            continue
        elif blockquote_text:
            # End of blockquote, add it
            quote_content = ' '.join(blockquote_text)
            quote_content = clean_markdown(quote_content)
            p = doc.add_paragraph(quote_content, style='Block Quote')
            blockquote_text = []

        # Handle headings
        if line.startswith('# ') and not line.startswith('## '):
            # Title (H1)
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            # Section heading (H2)
            heading = line[3:].strip()
            doc.add_paragraph(heading, style='Heading 2')
            i += 1
            continue

        if line.startswith('### '):
            # Subsection heading (H3)
            heading = line[4:].strip()
            doc.add_paragraph(heading, style='Heading 3')
            i += 1
            continue

        # Handle keyword line
        if line.startswith('**Keywords**:'):
            text = line.replace('**Keywords**:', 'Keywords:').strip()
            text = clean_markdown(text)
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run('Keywords: ')
            run.italic = True
            p.add_run(text.replace('Keywords:', '').strip())
            i += 1
            continue

        # Handle bullet points
        if line.strip().startswith('- **') or line.strip().startswith('- '):
            text = line.strip()[2:].strip()
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            i += 1
            continue

        # Handle numbered lists
        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            i += 1
            continue

        # Handle reference entries (lines starting with [number])
        if re.match(r'^\[\d+\]', line.strip()):
            text = clean_markdown(line.strip())
            p = doc.add_paragraph(text, style='Reference')
            i += 1
            continue

        # Handle tables - create actual Word table
        if line.strip().startswith('|'):
            # Collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                # Skip separator rows
                if '---' not in row_line:
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                i += 1

            # Create table if we have data
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'

                for row_idx, row_data in enumerate(table_rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = clean_markdown(cell_text)
                        # Format cell text
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.space_before = Pt(0)
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
                                # Bold header row
                                if row_idx == 0:
                                    run.font.bold = True

                # Add space after table
                doc.add_paragraph()
            continue

        # Regular paragraph
        if line.strip():
            text = clean_markdown(line)
            p = doc.add_paragraph(text)
            i += 1
            continue

        # Empty line
        i += 1

    # Handle any remaining blockquote
    if blockquote_text:
        quote_content = ' '.join(blockquote_text)
        quote_content = clean_markdown(quote_content)
        doc.add_paragraph(quote_content, style='Block Quote')

    # Save
    doc.save(output_path)
    print(f"Saved to {output_path}")

def clean_markdown(text):
    """Remove markdown formatting."""
    # Bold
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # Italic
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # Links [text](url)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    # Inline code
    text = re.sub(r'`([^`]+)`', r'\1', text)
    return text

if __name__ == '__main__':
    base_dir = Path(__file__).parent
    md_path = base_dir / 'paper.md'
    output_path = base_dir / 'exports' / 'paper_submission.docx'

    # Ensure exports directory exists
    output_path.parent.mkdir(exist_ok=True)

    create_formatted_docx(md_path, output_path)
