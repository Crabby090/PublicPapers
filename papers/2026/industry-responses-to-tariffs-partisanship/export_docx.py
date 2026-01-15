"""
Export paper.md to formatted DOCX for journal submission.
Times New Roman, double spacing, proper heading styles.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def create_formatted_docx(md_path, output_path):
    """Convert markdown to formatted DOCX."""
    content = Path(md_path).read_text(encoding='utf-8')

    doc = Document()
    styles = doc.styles

    # Normal style - Times New Roman 12pt, double spaced
    normal_style = styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    normal_style.paragraph_format.space_after = Pt(0)
    normal_style.paragraph_format.first_line_indent = Inches(0.5)

    # Heading 1 - Title
    h1_style = styles['Heading 1']
    h1_style.font.name = 'Times New Roman'
    h1_style.font.size = Pt(14)
    h1_style.font.bold = True
    h1_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h1_style.paragraph_format.space_before = Pt(0)
    h1_style.paragraph_format.space_after = Pt(0)
    h1_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 2 - Section headers
    h2_style = styles['Heading 2']
    h2_style.font.name = 'Times New Roman'
    h2_style.font.size = Pt(12)
    h2_style.font.bold = True
    h2_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h2_style.paragraph_format.space_before = Pt(12)
    h2_style.paragraph_format.space_after = Pt(0)
    h2_style.paragraph_format.first_line_indent = Inches(0)

    # Heading 3 - Subsection headers
    h3_style = styles['Heading 3']
    h3_style.font.name = 'Times New Roman'
    h3_style.font.size = Pt(12)
    h3_style.font.bold = True
    h3_style.font.italic = True
    h3_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    h3_style.paragraph_format.space_before = Pt(12)
    h3_style.paragraph_format.space_after = Pt(0)
    h3_style.paragraph_format.first_line_indent = Inches(0)

    # Quote style
    try:
        quote_style = styles.add_style('Block Quote', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        quote_style = styles['Block Quote']
    quote_style.font.name = 'Times New Roman'
    quote_style.font.size = Pt(11)
    quote_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    quote_style.paragraph_format.left_indent = Inches(0.5)
    quote_style.paragraph_format.right_indent = Inches(0.5)
    quote_style.paragraph_format.first_line_indent = Inches(0)

    # Reference style (hanging indent)
    try:
        ref_style = styles.add_style('Reference', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        ref_style = styles['Reference']
    ref_style.font.name = 'Times New Roman'
    ref_style.font.size = Pt(12)
    ref_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    ref_style.paragraph_format.left_indent = Inches(0.5)
    ref_style.paragraph_format.first_line_indent = Inches(-0.5)
    ref_style.paragraph_format.space_after = Pt(0)

    # Table caption style
    try:
        caption_style = styles.add_style('Table Caption', WD_STYLE_TYPE.PARAGRAPH)
    except ValueError:
        caption_style = styles['Table Caption']
    caption_style.font.name = 'Times New Roman'
    caption_style.font.size = Pt(12)
    caption_style.font.bold = True
    caption_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    caption_style.paragraph_format.space_before = Pt(6)
    caption_style.paragraph_format.space_after = Pt(0)
    caption_style.paragraph_format.first_line_indent = Inches(0)

    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    lines = content.split('\n')
    i = 0
    blockquote_text = []

    def clean_markdown(text):
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'\*(.+?)\*', r'\1', text)
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        return text

    while i < len(lines):
        line = lines[i]

        if line.strip() == '---':
            i += 1
            continue

        if line.startswith('>'):
            blockquote_text.append(line[1:].strip())
            i += 1
            continue
        elif blockquote_text:
            quote_content = clean_markdown(' '.join(blockquote_text))
            doc.add_paragraph(quote_content, style='Block Quote')
            blockquote_text = []

        if line.startswith('# ') and not line.startswith('## '):
            title = line[2:].strip()
            p = doc.add_paragraph(title, style='Heading 1')
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if line.startswith('## '):
            heading = line[3:].strip()
            doc.add_paragraph(heading, style='Heading 2')
            i += 1
            continue

        if line.startswith('### '):
            heading = line[4:].strip()
            doc.add_paragraph(heading, style='Heading 3')
            i += 1
            continue

        # Author/affiliation/date lines
        if line.startswith('**Author**:') or line.startswith('**Affiliation**:') or line.startswith('**Date**:'):
            text = clean_markdown(line).strip()
            p = doc.add_paragraph(text)
            p.paragraph_format.first_line_indent = Inches(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Keywords and JEL codes
        if line.startswith('**Keywords**:') or line.startswith('**JEL Codes**:'):
            label, value = line.split(':', 1)
            label = clean_markdown(label).replace('**', '').strip()
            value = clean_markdown(value).strip()
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Inches(0)
            run = p.add_run(f"{label}: ")
            run.italic = True
            p.add_run(value)
            i += 1
            continue

        # Table captions
        if line.strip().startswith('**Table '):
            text = clean_markdown(line).strip()
            p = doc.add_paragraph(text, style='Table Caption')
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1
            continue

        if line.strip().startswith('- '):
            text = clean_markdown(line.strip()[2:].strip())
            p = doc.add_paragraph(text, style='List Bullet')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            p.style.font.name = 'Times New Roman'
            p.style.font.size = Pt(12)
            i += 1
            continue

        if re.match(r'^\d+\.', line.strip()):
            text = re.sub(r'^\d+\.\s*', '', line.strip())
            text = clean_markdown(text)
            p = doc.add_paragraph(text, style='List Number')
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
            i += 1
            continue

        if re.match(r'^\[\d+\]', line.strip()):
            text = clean_markdown(line.strip())
            doc.add_paragraph(text, style='Reference')
            i += 1
            continue

        if line.strip().startswith('|'):
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if '---' not in row_line:
                    cells = [c.strip() for c in row_line.split('|')[1:-1]]
                    if cells:
                        table_rows.append(cells)
                i += 1
            if table_rows:
                num_cols = len(table_rows[0])
                table = doc.add_table(rows=len(table_rows), cols=num_cols)
                table.style = 'Table Grid'
                for row_idx, row_data in enumerate(table_rows):
                    row = table.rows[row_idx]
                    for col_idx, cell_text in enumerate(row_data):
                        cell = row.cells[col_idx]
                        cell.text = clean_markdown(cell_text)
                        for paragraph in cell.paragraphs:
                            paragraph.paragraph_format.space_after = Pt(0)
                            paragraph.paragraph_format.space_before = Pt(0)
                            for run in paragraph.runs:
                                run.font.name = 'Times New Roman'
                                run.font.size = Pt(10)
                                if row_idx == 0:
                                    run.font.bold = True
                doc.add_paragraph()
            continue

        if line.strip():
            text = clean_markdown(line)
            doc.add_paragraph(text)
            i += 1
            continue

        i += 1

    if blockquote_text:
        quote_content = clean_markdown(' '.join(blockquote_text))
        doc.add_paragraph(quote_content, style='Block Quote')

    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    print(f"Saved to {output_path}")


if __name__ == '__main__':
    base_dir = Path(__file__).parent
    md_path = base_dir / 'paper.md'
    output_path = base_dir / 'exports' / 'paper_submission.docx'
    create_formatted_docx(md_path, output_path)
